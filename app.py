import streamlit as st
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
import openai
import io
from collections import Counter
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer

# ------------------ CONFIG ------------------
st.set_page_config(page_title="SEO Content Wizard", page_icon="🔍")

# OpenAI API Key (Replace with your actual key)
OPENAI_API_KEY = "your-openai-api-key"

# Load NLP Model
nlp = spacy.load("en_core_web_sm")

# ------------------ FUNCTION TO CHECK OPENAI API ------------------
def check_openai_api():
    """Verifies OpenAI API key by making a test request."""
    openai.api_key = OPENAI_API_KEY
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": "Hello, test OpenAI response"}]
        )
        return True  # API is working
    except Exception as e:
        return str(e)  # Return error message if API fails

# ------------------ HEADER ------------------
st.title("🔍 SEO Content Wizard")
st.subheader("Analyze. Optimize. Rank Higher.")
st.write("Upload competitor URLs to analyze SEO factors and get AI-driven content recommendations.")

# ------------------ OPENAI API CHECK ------------------
api_status = check_openai_api()
if api_status is not True:
    st.error(f"⚠ OpenAI API Error: {api_status}")
    st.stop()  # Stop the program if API is not working

# ------------------ USER INPUT ------------------
keyword = st.text_input("Enter Target Keyword:")
urls = st.text_area("Enter 3-10 Competitor URLs (one per line)").split("\n")
urls = [url.strip() for url in urls if url.strip()]

if st.button("Analyze SEO"):
    if len(urls) < 3:
        st.error("Please enter at least 3 competitor URLs.")
    else:
        st.success("Analyzing URLs... Please wait.")

        # ------------------ SCRAPING FUNCTION ------------------
        results = []
        all_texts = []
        
        for url in urls:
            try:
                response = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
                soup = BeautifulSoup(response.text, "html.parser")

                title = soup.title.text if soup.title else "No Title"
                meta_desc = soup.find("meta", attrs={"name": "description"})
                meta_desc = meta_desc["content"] if meta_desc else "No Meta Description"
                word_count = len(soup.get_text().split())

                headings = [h.text.strip() for h in soup.find_all(["h1", "h2", "h3"])]
                internal_links = [a["href"] for a in soup.find_all("a", href=True) if url in a["href"]]
                external_links = [a["href"] for a in soup.find_all("a", href=True) if url not in a["href"]]

                text_content = soup.get_text()
                all_texts.append(text_content)

                results.append({
                    "URL": url,
                    "Meta Title": title,
                    "Meta Description": meta_desc,
                    "Word Count": word_count,
                    "H1-H3 Headings": headings,
                    "Internal Links": len(internal_links),
                    "External Links": len(external_links),
                    "Content": text_content
                })
            
            except Exception as e:
                st.error(f"Failed to analyze {url}: {str(e)}")

        df = pd.DataFrame(results)
        st.dataframe(df)

        # ------------------ NLP KEYWORD CLUSTERING ------------------
        def extract_keywords(text):
            doc = nlp(text.lower())
            words = [token.text for token in doc if token.is_alpha and not token.is_stop]
            return words

        all_keywords = []
        for text in all_texts:
            all_keywords.extend(extract_keywords(text))

        keyword_freq = Counter(all_keywords)
        st.subheader("🔑 Keyword Clustering & Frequency")
        st.write(pd.DataFrame(keyword_freq.most_common(20), columns=["Keyword", "Frequency"]))

        # ------------------ COMPETITOR CONTENT GAP ANALYSIS ------------------
        tfidf = TfidfVectorizer(stop_words="english")
        tfidf_matrix = tfidf.fit_transform(all_texts)
        feature_names = tfidf.get_feature_names_out()

        keyword_importance = {feature_names[i]: tfidf_matrix.sum(axis=0).tolist()[0][i] for i in range(len(feature_names))}
        sorted_keywords = sorted(keyword_importance.items(), key=lambda x: x[1], reverse=True)

        st.subheader("📊 Competitor Content Gap Analysis")
        st.write(pd.DataFrame(sorted_keywords[:20], columns=["Keyword", "Relevance Score"]))

        # ------------------ AI-POWERED RECOMMENDATIONS ------------------
        st.subheader("🧠 AI-Powered SEO Recommendations")
        prompt = f"""
        Based on competitor content analysis, generate an SEO-optimized content outline for the keyword '{keyword}'.

        Competitor Data:
        {df.to_string()}
        
        Include:
        - Optimized meta title & description
        - Ideal word count
        - Recommended H1, H2, and H3 headings
        - Missing subtopics to enhance SEO
        """
        
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "system", "content": "You are an SEO expert."}, {"role": "user", "content": prompt}]
            )
            ai_recommendations = response["choices"][0]["message"]["content"]
            st.text_area("SEO Content Recommendations:", ai_recommendations, height=250)
        except Exception as e:
            st.error(f"OpenAI Error: {str(e)}")

        # ------------------ EXPORT OPTIONS ------------------
        def export_to_csv(data):
            return df.to_csv(index=False).encode('utf-8')

        def export_to_docx(data):
            doc = Document()
            doc.add_heading("SEO Content Analysis", 0)
            for row in data:
                doc.add_paragraph(f"🔗 URL: {row['URL']}")
                doc.add_paragraph(f"📌 Meta Title: {row['Meta Title']}")
                doc.add_paragraph(f"📝 Meta Description: {row['Meta Description']}")
                doc.add_paragraph(f"📄 Word Count: {row['Word Count']}")
                doc.add_paragraph(f"📌 Headings: {', '.join(row['H1-H3 Headings'])}")
                doc.add_paragraph(f"🔗 Internal Links: {row['Internal Links']}")
                doc.add_paragraph(f"🔗 External Links: {row['External Links']}")
                doc.add_paragraph("-" * 50)
            buffer = io.BytesIO()
            doc.save(buffer)
            return buffer.getvalue()

        def export_to_txt(data):
            text = "SEO Content Analysis Report\n\n"
            for row in data:
                text += f"URL: {row['URL']}\n"
                text += f"Meta Title: {row['Meta Title']}\n"
                text += f"Meta Description: {row['Meta Description']}\n"
                text += f"Word Count: {row['Word Count']}\n"
                text += "-" * 50 + "\n"
            return text.encode('utf-8')

        col1, col2, col3 = st.columns(3)
        with col1:
            st.download_button(label="📜 Download CSV", data=export_to_csv(results), file_name="SEO_Analysis.csv", mime="text/csv")
        with col2:
            st.download_button(label="📄 Download DOCX", data=export_to_docx(results), file_name="SEO_Analysis.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with col3:
            st.download_button(label="📄 Download TXT", data=export_to_txt(results), file_name="SEO_Analysis.txt", mime="text/plain")
